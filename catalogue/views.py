# -*- coding: utf-8 -*-
#python
from __future__ import unicode_literals
import docx
from docx.shared import Inches

#django
from django.utils import timezone
from django.http import JsonResponse, StreamingHttpResponse, HttpResponse, FileResponse
from django.utils.six import StringIO

#rest
from rest_framework.views import APIView
from rest_framework.pagination import PageNumberPagination

#lib
from education.lib.baseviews import BaseApiView

#apps
from education.catalogue.models import Category, Example
from education.catalogue.serializers import (CategorySerializer,
    ExampleSerializer, ExampleDetailSerializer)
from education.catalogue.forms import (CategoryCreateForm, CategoryUpdateForm,
    ExampleCreateForm, ExampleUpdateForm)
# Create your views here.
################################################################################
#                              分页                                            #
################################################################################

class Page(PageNumberPagination):
    # 每页显示的数据条数
    max_page_size = 30
    page_size = 30
    page_size_query_param = 'size'
    # 页码
    page_query_param = 'page'

################################################################################
#                              分类                                            #
################################################################################

# 分类创建及更新
class UpdateCategoryView(BaseApiView):
    def post(self, request, *args, **kw):
        if kw.get('pk'):
            form = CategoryUpdateForm(request.data,
                request.FILES,
                kw.get('pk'))
        else:
            form = CategoryCreateForm(request.data,
                request.FILES,
                kw.get('parent_pk'))

        if form.is_valid():
            res = {
                'msg': 'success',
                'desc': 'success',
                'code': 0,
            }

            model = form.save()
            res['data'] = {'id': model.id}
        else:
            res = self.err_response(form)
        return JsonResponse(res)

#获取分类列表
class CategoryView(APIView):
    def get(self, request, *arg, **kw):
        parent_pk = kw.get('parent_pk')
        res = {
            'msg': 'success',
            'desc': 'success',
            'code': 0,
        }
        if not parent_pk:
            tree = Category.get_root_nodes()
            data = CategorySerializer(tree, many=True)
            res['data'] = data.data
            return JsonResponse(res, safe=False)
        try:
            node = Category.objects.get(id=parent_pk)
            data = CategorySerializer(node.get_children(), many=True)
        except Category.DoesNotExist:
            return JsonResponse({'msg':'invaild id',
                'desc': '该分类不存在'}, status=404)
        else:
            res['data'] = data.data
            return JsonResponse(res, safe=False)

# 或许父节点分类树
class AncestorsCategoryView(APIView):
    def get(self, request, *arg, **kw):
        category_id = kw.get('category_pk')
        res = {
            'msg': 'success',
            'desc': 'success',
            'code': 0,
        }
        node = Category.objects.get(id=category_id)
        category_qs = list(node.get_ancestors()) + [node]
        data = CategorySerializer(category_qs, many=True)
        res['data'] = data.data
        return JsonResponse(res, safe=False)

################################################################################
#                              题目                                            #
################################################################################

# 题目创建及更新
class UpdateExampleView(BaseApiView):
    def post(self, request, *args, **kw):
        print request.data
        if kw.get('pk'):
            form = ExampleUpdateForm(request.data,
                kw.get('pk'), request.FILES,)
        else:
            form = ExampleCreateForm(request.data,
                request.FILES,)
        if form.is_valid():
            res = {
                'msg': 'success',
                'desc': 'success',
                'code': 0,
            }
            example = form.save()
            data = ExampleDetailSerializer(example)
            res['data'] = data.data
            # res['data'] = UserInfoSerializer(form.user).data
        else:
            res = self.err_response(form)
        return JsonResponse(res)

# 获取题目列表
class ExampleView(BaseApiView):
    def get(self, request, *args, **kw):
        res = {
            'msg': 'success',
            'desc': 'success',
            'code': 0,
        }
        category_pk = kw.get('category_pk')
        if category_pk:
            manager = Category.objects.get(id=category_pk).examples
        else:
            manager = Example.objects

        # 用于之后过滤数据 暂时没用
        example_qs = manager.all()
        page = Page()
        page_data = page.paginate_queryset( \
            queryset=example_qs, request=request, view=self)
        data = ExampleSerializer(page_data, many=True)
        res['data'] = {'examples': data.data,
            'next_link': page.get_next_link()}
        return JsonResponse(res, safe=False)

class ExampleDetaiView(APIView):
    def get(self, request, *args, **kw):
        res = {
            'msg': 'success',
            'desc': 'success',
            'code': 0,
        }
        example_pk = kw.get('pk')
        example = Example.objects.get(id=example_pk)

        # 用于之后过滤数据 暂时没用
        data = ExampleDetailSerializer(example)
        res['data'] = data.data
        return JsonResponse(res, safe=False)


class DocxView(BaseApiView):
    file_word = docx.Document()
    IObuf = StringIO()
    def get(self, request, *args, **kw):
        example_id_list = request.GET.get('example_ids', '').split('-')
        examples = Example.objects.filter(id__in=example_id_list)
        # examples = Example.objects.all()
        self.examples = examples
        self.write_heading()
        self.write_example()
        self.file_word.add_section()
        self.write_answer_heading()
        self.write_answer()
        zbuf = StringIO()
        self.file_word.save(zbuf)
        response = HttpResponse(zbuf.getvalue(),content_type='application/msword')
        response['Content-Disposition'] = 'attachment;filename=%s.doc' \
            % self.file_name
        return response

    def get_file_name(self):
        self.file_name = timezone.now().strftime('%Y-%m-%d %H:%M:%S ') + 'examples'
        return self.file_name

    def write_heading(self):
        head = self.get_file_name()
        self.file_word.add_heading(head, 0)

    def write_answer_heading(self):
        self.file_word.add_heading('答案', 0)

    def write_example(self):
        for index, e in enumerate(self.examples):
            self.file_word.add_paragraph(''.join((str(index + 1), ': ',  e.content)))
            self.write_images(e.images.all())
            # 一题空1行
            self.file_word.add_paragraph('')

    def write_answer(self):
        for index, e in enumerate(self.examples):
            self.file_word.add_paragraph(''.join(('题', str(index + 1), ': ')))
            self.write_answer_of_example(e.answers.all())
            # 一题空1行
            self.file_word.add_paragraph('')

    def write_answer_of_example(self, answer_qs):
        for index, a in enumerate(answer_qs):
            self.file_word.add_paragraph(''.join(('　　解法', str(index + 1),
                ': ',
                a.answer)))
            self.write_images(a.images.all())

    def write_images(self, images_qs):
        if not images_qs:
            return
        # pic_name_list = [' ' * 7, ]
        for i, img in enumerate(images_qs):
            #　一行四图
            if (i+1) % 4 == 1:
                pic_name_list = [' ' * 7, ] # 首行全角空三格对齐图片
                pic_paragraph = self.file_word.add_paragraph().add_run()
            self.IObuf.write(img.image.read())
            pic = pic_paragraph.add_picture(self.IObuf,
                width=Inches(1.0),
                height=Inches(1.0))
            self.IObuf.truncate()

            # XXX:
            pic_name_list.append(('图' + str(i+1)))
            # 四图后换行插入每个图片名称
            if (i+1) % 4 == 0:
                pic_paragraph = self.file_word.add_paragraph()
                for img_name in pic_name_list:
                    pic_paragraph.add_run(img_name)
            else:
                pic_name_list.append(' ' * 13) # 每张图片名称间距七个对齐图片
        if len(pic_name_list) != 1:
            pic_paragraph = self.file_word.add_paragraph()
            for img_name in pic_name_list:
                pic_paragraph.add_run(img_name)